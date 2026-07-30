"""Multi-format document parser: PDF, DOCX, CSV, TXT, MD, HTML."""

import csv
import io
from pathlib import Path
from typing import Iterator

import fitz  # PyMuPDF
import markdown
from bs4 import BeautifulSoup
from docx import Document as DocxDocument


class ParsedDocument:
    """Parsed document with extracted text and metadata."""

    def __init__(self, text: str, metadata: dict):
        self.text = text
        self.metadata = metadata  # page_number, section_title, etc.


def parse_pdf(file_path: str) -> list[ParsedDocument]:
    """Parse a PDF file, returning one ParsedDocument per page."""
    docs = []
    with fitz.open(file_path) as pdf:
        for page_num, page in enumerate(pdf, start=1):
            text = page.get_text("text")
            if text.strip():
                docs.append(
                    ParsedDocument(
                        text=text.strip(),
                        metadata={"page_number": page_num},
                    )
                )
    return docs


def parse_docx(file_path: str) -> list[ParsedDocument]:
    """Parse a DOCX file, splitting by headings as sections."""
    doc = DocxDocument(file_path)
    docs = []
    current_text: list[str] = []
    current_section = "正文"

    for para in doc.paragraphs:
        if para.style.name.startswith("Heading"):
            # Save previous section
            if current_text:
                docs.append(
                    ParsedDocument(
                        text="\n".join(current_text).strip(),
                        metadata={"section_title": current_section},
                    )
                )
                current_text = []
            current_section = para.text.strip() or para.style.name
        else:
            if para.text.strip():
                current_text.append(para.text)

    # Don't forget the last section
    if current_text:
        docs.append(
            ParsedDocument(
                text="\n".join(current_text).strip(),
                metadata={"section_title": current_section},
            )
        )

    # Also extract tables
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(" | ".join(cells))
        if rows:
            docs.append(
                ParsedDocument(
                    text="\n".join(rows),
                    metadata={"section_title": "表格"},
                )
            )

    return docs


def parse_csv(file_path: str) -> list[ParsedDocument]:
    """Parse a CSV file, treating each row as a product record."""
    docs = []
    with open(file_path, "r", encoding="utf-8-sig") as f:
        reader = csv.DictReader(f)
        if not reader.fieldnames:
            # Try without header
            f.seek(0)
            reader = csv.reader(f)
            for i, row in enumerate(reader):
                text = " | ".join(str(cell) for cell in row)
                if text.strip():
                    docs.append(
                        ParsedDocument(
                            text=text,
                            metadata={"row_number": i + 1},
                        )
                    )
            return docs

        for i, row in enumerate(reader):
            parts = [f"{key}: {value}" for key, value in row.items() if value]
            text = "\n".join(parts)
            if text.strip():
                docs.append(
                    ParsedDocument(
                        text=text,
                        metadata={
                            "row_number": i + 2,  # +2 for header
                            **{k: v for k, v in row.items() if v},
                        },
                    )
                )

    return docs


def parse_txt(file_path: str) -> list[ParsedDocument]:
    """Parse a plain text file."""
    with open(file_path, "r", encoding="utf-8") as f:
        text = f.read().strip()

    if not text:
        return []

    # Split by double newlines as paragraph markers
    paragraphs = text.split("\n\n")
    docs = []
    for i, para in enumerate(paragraphs):
        if para.strip():
            docs.append(
                ParsedDocument(
                    text=para.strip(),
                    metadata={"paragraph_index": i},
                )
            )

    return docs


def parse_markdown(file_path: str) -> list[ParsedDocument]:
    """Parse a Markdown file, extracting text content."""
    with open(file_path, "r", encoding="utf-8") as f:
        md_text = f.read()

    # Convert to plain text via HTML intermediate
    html = markdown.markdown(md_text)
    soup = BeautifulSoup(html, "html.parser")
    text = soup.get_text(separator="\n", strip=True)

    if not text:
        return []

    # Split by double newlines as sections
    sections = text.split("\n\n")
    docs = []
    current_section = ""
    for section in sections:
        section = section.strip()
        if not section:
            continue
        # Check if this looks like a heading
        if section.startswith("#") or len(section) < 80:
            current_section = section
        docs.append(
            ParsedDocument(
                text=section,
                metadata={
                    "section_title": current_section if current_section else None,
                },
            )
        )

    return docs


def parse_html(file_path: str) -> list[ParsedDocument]:
    """Parse an HTML file, extracting text with heading hierarchy."""
    with open(file_path, "r", encoding="utf-8") as f:
        html_text = f.read()

    soup = BeautifulSoup(html_text, "html.parser")

    # Remove script and style elements
    for tag in soup(["script", "style", "nav", "footer", "header"]):
        tag.decompose()

    # Extract by headings as sections
    docs = []
    current_heading = "正文"
    current_text: list[str] = []

    for element in soup.find_all(["h1", "h2", "h3", "h4", "h5", "h6", "p", "table", "li"]):
        if element.name.startswith("h"):
            if current_text:
                docs.append(
                    ParsedDocument(
                        text="\n".join(current_text).strip(),
                        metadata={"section_title": current_heading},
                    )
                )
                current_text = []
            current_heading = element.get_text(strip=True)
        elif element.name == "table":
            rows = []
            for row in element.find_all("tr"):
                cells = [cell.get_text(strip=True) for cell in row.find_all(["td", "th"])]
                rows.append(" | ".join(cells))
            if rows:
                current_text.append("\n".join(rows))
        else:
            text = element.get_text(strip=True)
            if text:
                current_text.append(text)

    if current_text:
        docs.append(
            ParsedDocument(
                text="\n".join(current_text).strip(),
                metadata={"section_title": current_heading},
            )
        )

    return docs


PARSERS = {
    "pdf": parse_pdf,
    "docx": parse_docx,
    "csv": parse_csv,
    "txt": parse_txt,
    "md": parse_markdown,
    "html": parse_html,
}


def parse_document(file_path: str, file_type: str) -> list[ParsedDocument]:
    """Parse a document using the appropriate parser based on file type."""
    parser = PARSERS.get(file_type.lower())
    if not parser:
        raise ValueError(f"Unsupported file type: {file_type}")

    return parser(file_path)
